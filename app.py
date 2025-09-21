from flask import Flask, render_template, request, send_file
from docx import Document
import io

app = Flask(__name__)

def get_style_mapping(format_doc):
    style_mapping = {}
    for para in format_doc.paragraphs:
        name = para.style.name
        if name not in style_mapping:
            style_mapping[name] = para.style
    for table in format_doc.tables:
        style_mapping['Table'] = table.style
    return style_mapping

def apply_formatting(input_doc, format_doc):
    input_document = Document(input_doc)
    format_document = Document(format_doc)
    style_mapping = get_style_mapping(format_document)

    for para in input_document.paragraphs:
        if para.style.name.startswith('Heading'):
            heading_level = para.style.name
            if heading_level in style_mapping:
                para.style = style_mapping[heading_level]
        else:
            if 'Normal' in style_mapping:
                para.style = style_mapping['Normal']

    for table in input_document.tables:
        if 'Table' in style_mapping:
            table.style = style_mapping['Table']

    output = io.BytesIO()
    input_document.save(output)
    output.seek(0)
    return output

@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        input_file = request.files.get('input_file')
        format_file = request.files.get('format_file')

        if not input_file or not format_file:
            return "Both files are required!", 400

        output = apply_formatting(input_file, format_file)
        return send_file(
            output,
            as_attachment=True,
            download_name='converted.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    return render_template('index.html')

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)
